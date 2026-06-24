"""
Assemble the final thesis .docx.

Strategy: start from the pandoc-converted body (which already has proper
Heading 1/2/3 styles), enforce DOE body formatting, then PREPEND the
official DOE title page and a Table-of-Contents field. Working from the
styled body keeps every heading style intact, so the TOC field populates
correctly in Word.

Output: thesis/Final_Thesis_Khush_Patel.docx
"""
from pathlib import Path
import copy
from docx import Document
from docx.shared import Pt, Mm, Cm, RGBColor
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
HERE = Path(__file__).resolve().parent

TITLE_TEMPLATE = ROOT / "format" / "BS_Thesis_title_page_2026.docx"
BODY = HERE / "body.docx"
OUT = HERE / "Final_Thesis_Khush_Patel.docx"

TOPIC = ("When Does Memory Help an LLM Judge? "
         "Evaluation of Memory-Augmented Judging")
STUDENT = "Khush Patel"
SUPERVISOR = "Prof. Bader Rasheed"
FONT = "Times New Roman"


def fill_cell(cell, text, *, size=14, bold=True):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.name = FONT
    r.font.size = Pt(size)


def force_font(style):
    """DOE body style: TNR 14pt, black, 1.5 spacing, 1.25cm first-line indent."""
    style.font.name = FONT
    style.font.size = Pt(14)
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.first_line_indent = Cm(1.25)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT)


def set_margins(doc):
    """DOE margins: left 25 mm, right/top/bottom 20 mm."""
    for s in doc.sections:
        s.left_margin = Mm(25)
        s.right_margin = Mm(20)
        s.top_margin = Mm(20)
        s.bottom_margin = Mm(20)


def remove_consultant_row(sig_tbl):
    """DOE rule 8: if there is no consultant, delete the consultant row."""
    for row in list(sig_tbl.rows):
        if "Консультант" in row.cells[0].text or "Consultant" in row.cells[0].text:
            row._tr.getparent().remove(row._tr)


def blacken_links(doc):
    """DOE rule 5: hyperlink / reference text must be black, not blue."""
    for run in doc.element.iter(qn("w:r")):
        rPr = run.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            run.insert(0, rPr)
        color = rPr.find(qn("w:color"))
        if color is None:
            color = OxmlElement("w:color")
            rPr.append(color)
        color.set(qn("w:val"), "000000")


def add_page_numbers(doc):
    """DOE rule 7: continuous page numbers, centred in the footer.
    The title page counts as page 1 but shows no number, so the first
    section is given a 'different first page' footer left blank."""
    for section in doc.sections:
        section.different_first_page_header_footer = True
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end")
        run._r.append(fb)
        run._r.append(instr)
        run._r.append(fe)
        # first-page footer left empty -> no number on the title page
        section.first_page_footer.is_linked_to_previous = False


def renumber_cross_references(doc):
    """Rewrite body cross-references to match the sequential captions.

    Pandoc renders LaTeX \\ref{} as internal hyperlinks: the bare number
    ("4.1") sits inside a <w:hyperlink>, while the word "Figure"/"Table" is
    an ordinary run just before it. Each hyperlinked number is mapped to its
    sequential value; the preceding text disambiguates Figure vs Table."""
    fig = {"4.1": "1", "4.2": "2", "4.3": "3"}
    tab = {"4.1": "1", "4.2": "2", "4.3": "3", "4.4": "4",
           "7.1": "5", "7.2": "6", "8.1": "7"}
    count = 0
    for p in doc.paragraphs:
        children = list(p._p)
        for idx, el in enumerate(children):
            if el.tag != qn("w:hyperlink"):
                continue
            tnodes = el.findall(".//" + qn("w:t"))
            cur = "".join(t.text or "" for t in tnodes).strip()
            if cur not in fig and cur not in tab:
                continue
            context = ""
            for back in range(idx - 1, max(-1, idx - 4), -1):
                bt = "".join(t.text or "" for t in children[back].iter(qn("w:t")))
                context = bt + context
                if "Figure" in context or "Table" in context:
                    break
            new = None
            if "Figure" in context and cur in fig:
                new = fig[cur]
            elif "Table" in context and cur in tab:
                new = tab[cur]
            if new is not None and tnodes:
                tnodes[0].text = new
                for t in tnodes[1:]:
                    t.text = ""
                count += 1
    return count

def number_captions(doc):
    """DOE rules 9 & 10: tables get 'Table N – ' above (centred), figures get
    'Figure N – ' below (left-aligned). Pandoc drops the numbers, so they are
    added here in document order."""
    tbl_n = 0
    fig_n = 0
    for p in doc.paragraphs:
        sty = p.style.name
        if sty == "Table Caption":
            tbl_n += 1
            txt = p.text.strip()
            if not txt.lower().startswith("table"):
                _prefix_caption(p, f"Table {tbl_n} – ")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif sty == "Image Caption":
            fig_n += 1
            txt = p.text.strip()
            if not txt.lower().startswith("figure"):
                _prefix_caption(p, f"Figure {fig_n} – ")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return tbl_n, fig_n


def number_headings(doc):
    """Add chapter/section numbers to headings so the generated Table of
    Contents reads '1 Introduction', '2.1 History', '2.1.1 ...'.

    The Abstract is front matter and stays unnumbered. The four appendices
    are lettered A-D. Numbering is applied by writing the number into the
    first run of each heading paragraph."""
    ch = 0          # chapter counter
    sec = 0         # section counter
    sub = 0         # subsection counter
    app = 0         # appendix counter (A, B, C, D)
    in_appendix = False
    APPENDIX_TITLES = {
        "Frozen-Memory Audit Record Format",
        "Evaluation Modes and Memory Conditions",
        "Full Per-Mode Results",
        "Reproduction",
    }
    FRONT_MATTER = {"Abstract", "References", "Contents"}

    for p in doc.paragraphs:
        sty = p.style.name
        text = p.text.strip()
        if not text:
            continue

        if sty == "Heading 1":
            if text in FRONT_MATTER:
                continue
            if text in APPENDIX_TITLES:
                in_appendix = True
                app += 1
                _prefix_heading(p, f"Appendix {chr(64 + app)}. ")
                sec = sub = 0
            else:
                ch += 1
                sec = sub = 0
                _prefix_heading(p, f"{ch}. ")
        elif sty == "Heading 2" and ch > 0 and not in_appendix:
            sec += 1
            sub = 0
            _prefix_heading(p, f"{ch}.{sec}. ")
        elif sty == "Heading 3" and ch > 0 and not in_appendix:
            sub += 1
            _prefix_heading(p, f"{ch}.{sec}.{sub}. ")


def _prefix_heading(paragraph, prefix):
    """Insert a numbering prefix at the start of a heading paragraph."""
    new_r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = prefix
    new_r.append(t)
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is not None:
        pPr.addnext(new_r)
    else:
        paragraph._p.insert(0, new_r)


def _prefix_caption(paragraph, prefix):
    """Insert a bold prefix run at the start of a caption paragraph."""
    new_r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    b = OxmlElement("w:b")
    rPr.append(b)
    new_r.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = prefix
    new_r.append(t)
    paragraph._p.insert(
        list(paragraph._p).index(paragraph._p.find(qn("w:pPr"))) + 1
        if paragraph._p.find(qn("w:pPr")) is not None else 0,
        new_r,
    )


def make_para(text="", *, bold=False, italic=False, size=14,
              align=WD_ALIGN_PARAGRAPH.CENTER):
    """Build a standalone <w:p> element (centred, Times New Roman)."""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), {WD_ALIGN_PARAGRAPH.CENTER: "center",
                         WD_ALIGN_PARAGRAPH.LEFT: "left"}[align])
    pPr.append(jc)
    p.append(pPr)
    if text:
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rf = OxmlElement("w:rFonts")
        for a in ("w:ascii", "w:hAnsi", "w:cs"):
            rf.set(qn(a), FONT)
        rPr.append(rf)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(size * 2))
        rPr.append(sz)
        if bold:
            rPr.append(OxmlElement("w:b"))
        if italic:
            rPr.append(OxmlElement("w:i"))
        r.append(rPr)
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        r.append(t)
        p.append(r)
    return p


def make_pagebreak():
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)
    return p


def make_toc():
    """Caption + a TOC field that Word fills in on 'Update Field'."""
    elems = [make_para("Contents", bold=True, size=16),
             make_para("")]
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fs = OxmlElement("w:fldChar"); fs.set(qn("w:fldCharType"), "separate")
    ph = OxmlElement("w:t")
    ph.text = "Update this field in Word to generate the table of contents."
    fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end")
    for n in (fb, instr, fs, ph, fe):
        r.append(n)
    p.append(r)
    elems.append(p)
    return elems


def main():
    # --- start from the styled body -------------------------------------
    doc = Document(str(BODY))
    force_font(doc.styles["Normal"])
    set_margins(doc)                       # DOE rule 2: 25/20/20/20

    body_el = doc.element.body
    first = body_el[0]   # first existing element -> insert title page before it

    # --- title page (from the official template) ------------------------
    tpl = Document(str(TITLE_TEMPLATE))
    fill_cell(tpl.tables[0].cell(0, 2), TOPIC, size=14)
    fill_cell(tpl.tables[1].cell(0, 2), STUDENT, size=12)
    fill_cell(tpl.tables[1].cell(2, 2), SUPERVISOR, size=12)
    remove_consultant_row(tpl.tables[1])   # DOE rule 8: no consultant

    prepend = []
    for el in list(tpl.element.body):
        if el.tag == qn("w:sectPr"):
            continue
        prepend.append(copy.deepcopy(el))
    prepend.append(make_pagebreak())
    prepend.extend(make_toc())
    prepend.append(make_pagebreak())

    for el in prepend:
        first.addprevious(el)

    # --- "References" heading before the bibliography -------------------
    ref_inserted = False
    paras = doc.paragraphs
    model_pPr = None
    for p in paras:
        if p.text.strip() in ("Conclusion", "Introduction"):
            mp = p._p.find(qn("w:pPr"))
            if mp is not None:
                model_pPr = mp
            break
    for i, p in enumerate(paras):
        t = p.text.strip()
        is_ref = (len(t) > 25 and t[0].isalpha() and t[0].isupper()
                  and ", " in t[:45]
                  and any(f" {y}" in t for y in range(1990, 2031)))
        if is_ref and i + 1 < len(paras):
            nxt = paras[i + 1].text.strip()
            if nxt and ", " in nxt[:45]:
                hp = OxmlElement("w:p")
                if model_pPr is not None:
                    hp.append(copy.deepcopy(model_pPr))
                rr = OxmlElement("w:r")
                tt = OxmlElement("w:t")
                tt.text = "References"
                rr.append(tt)
                hp.append(rr)
                p._p.addprevious(hp)
                ref_inserted = True
                break

    # number the chapters/sections so the TOC reads "1 Introduction" etc.
    number_headings(doc)
    # DOE rule 5: force all run text (incl. references / links) to black
    blacken_links(doc)
    # DOE rule 7: continuous page numbers, none on the title page
    add_page_numbers(doc)
    # DOE rules 9 & 10: numbered table/figure captions
    tbl_n, fig_n = number_captions(doc)
    # keep body cross-references consistent with the sequential captions
    xref_n = renumber_cross_references(doc)

    doc.save(str(OUT))
    print(f"written: {OUT}")
    print(f"  margins 25/20/20/20; indent 1.25cm; links black; page numbers on")
    print(f"  consultant row removed; title page prepended; TOC field added")
    print(f"  captions numbered: {tbl_n} tables, {fig_n} figures")
    print(f"  references heading: {'inserted' if ref_inserted else 'NOT FOUND'}")


if __name__ == "__main__":
    main()
