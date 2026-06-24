"""
Build the Russian annotation as a .docx.

Uses the official DOE annotation title-page template for page 1, then
appends the condensed Russian annotation. DOE formatting: Times New
Roman 14pt, 1.5 spacing, margins 25/20/20/20, continuous page numbers
(requirement 11), first-line indent 1.25 cm on body paragraphs.

Output: thesis/Final_Annotation_Khush_Patel.docx
"""
from pathlib import Path
import copy
from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
HERE = Path(__file__).resolve().parent
TITLE_TEMPLATE = ROOT / "format" / "BS_Annotation_title_page_2026.docx"
OUT = HERE / "Final_Annotation_Khush_Patel.docx"

FONT = "Times New Roman"
TOPIC = ("Использование памяти и графов знаний для оценки "
         "действий интеллектуальных агентов")
STUDENT = "Пател Кхуш Прадипбхаи"

# --- annotation content: (level, text) ; level 0 = chapter heading,
#     1 = section heading, 2 = body paragraph ----------------------------
CONTENT = [
 (0, "Введение"),
 (2, "Большие языковые модели применяются не только для генерации текста, но и для его оценивания в парадигме «языковая модель в роли судьи». Стандартный судья однопроходный: он оценивает каждый пример изолированно и ничего не запоминает. Подключение к судье постоянной памяти потенциально позволяет переиспользовать прошлые оценки, однако та же память может распространять прошлые ошибки и приводить к утечке информации между связанными примерами. Цель работы — спроектировать, реализовать и строго оценить такую систему и определить, при каких условиях память действительно повышает качество оценивания. Научная новизна состоит в разработанном протоколе оценивания без утечек данных, в двухуровневом аудите замороженной памяти с криптографическим отпечатком SHA-256, и в корректирующем эмпирическом результате, опровергающем правдоподобный, но необоснованный вывод о пользе памяти."),

 (0, "Краткое описание основной части"),
 (2, "Проблема. Память, подключённая к судье, может улучшать оценивание, но также может порождать утечку данных между связанными примерами; неаудируемый стенд оценивания может сообщать как о несуществующем выигрыше, так и о несуществующем падении точности."),
 (2, "Решение. Система MCTS-MAJ состоит из судьи с памятью, поискового судьи на основе дерева Монте-Карло и модуля поискового извлечения; память хранится в графовой базе Neo4j с пятью типами узлов и типизированными связями. Разработан протокол оценивания без утечек данных, закрывающий два канала: разбиение данных по вопросам и замораживание памяти на время теста. Реализован двухуровневый аудит замороженной памяти: детерминированный снимок состояния «до и после» с отпечатком SHA-256 и перехват во время выполнения всех операций записи. Аудит уже на этапе разработки выявил реальный дефект утечки, который был исправлен. Точность сопровождается доверительным интервалом Уилсона; сравнения режимов выполняются на парных данных с точным критерием Макнемара и парным бутстрэп-интервалом."),
 (2, "Результаты. Эксперименты используют открытый набор данных EvalsBench (160 примеров, 80 уникальных вопросов, сбалансированное распределение проходных и непроходных ответов). При обычном протоколе память создавала видимость выигрыша порядка шести-десяти процентных пунктов. При аудируемом протоколе на 80 отложенных примерах для модели GPT-4o точность однопроходного судьи составила 70,0 %, судьи с памятью 65,0 %, поискового судьи 70,0 %, объединённого режима 71,8 %; все доверительные интервалы перекрываются, ни одно различие не значимо. Эталонная память с истинными метками не превосходит самостоятельно сформированную. «Отравление» памяти не вызывает значимого снижения точности; ранее наблюдавшийся обвал до 21,2 % оказался артефактом: сетевой сбой засчитывал ошибочные запросы как неверные ответы, и исправление обработки ошибок полностью устранило этот эффект."),

 (0, "Заключение"),
 (2, "В работе спроектирована и реализована система MCTS-MAJ; разработан протокол оценивания без утечек данных и двухуровневый аудит замороженной памяти. При обычном протоколе память создавала видимость значительного выигрыша; при аудируемом протоколе этот выигрыш не подтверждается. Кажущееся резкое падение точности при «отравлении» памяти оказалось артефактом измерения. Защищаемый вклад работы является методологическим: разбиение по вопросам, замороженная память и двухуровневый аудит образуют протокол, доказывающий отсутствие утечки данных. На данном наборе данных как кажущаяся польза, так и кажущаяся хрупкость памяти в значительной мере являются артефактами протокола оценивания. Направления дальнейшей работы: расширение набора данных, применение аудируемого протокола к более слабым и более сильным моделям-судьям, оформление протокола и аудита как самостоятельного инструмента."),

 (0, "Список использованной литературы"),
 (3, "L. Zheng et al., “Judging LLM-as-a-judge with MT-bench and chatbot arena,” NeurIPS Datasets and Benchmarks, 2023."),
 (3, "J. Gu et al., “A survey on LLM-as-a-judge,” arXiv:2411.15594, 2024."),
 (3, "P. Lewis et al., “Retrieval-augmented generation for knowledge-intensive NLP tasks,” NeurIPS, 2020."),
 (3, "Z. Zhang et al., “A survey on the memory mechanism of LLM-based agents,” arXiv:2404.13501, 2024."),
 (3, "W. Zou, R. Geng, B. Wang, and J. Jia, “PoisonedRAG,” USENIX Security, 2024."),
 (3, "S. Balloccu et al., “Leak, cheat, repeat: Data contamination and evaluation malpractices in closed-source LLMs,” EACL, 2024."),
]


def add_page_numbers(doc):
    """Continuous centred page numbers; title page (first page) has none."""
    for section in doc.sections:
        section.different_first_page_header_footer = True
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
        fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end")
        run._r.append(fb); run._r.append(instr); run._r.append(fe)


def ensure_heading_styles(doc):
    """Create Heading 1 / Heading 2 paragraph styles if the template lacks
    them, so a Word TOC field can detect and number the headings."""
    from docx.enum.style import WD_STYLE_TYPE
    styles = doc.styles
    existing = {s.name for s in styles}
    specs = [("Heading 1", 14, 1), ("Heading 2", 14, 2)]
    for name, size, outline in specs:
        if name in existing:
            st = styles[name]
        else:
            st = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        st.font.name = FONT
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = None
        pf = st.paragraph_format
        pf.first_line_indent = Cm(0)
        pf.space_before = Pt(6)
        pf.space_after = Pt(3)
        pf.keep_with_next = True
        # set the outline level so the TOC field picks it up
        pPr = st.element.get_or_add_pPr()
        ol = pPr.find(qn("w:outlineLvl"))
        if ol is None:
            ol = OxmlElement("w:outlineLvl")
            pPr.append(ol)
        ol.set(qn("w:val"), str(outline - 1))


def main():
    # --- title page from the official template --------------------------
    doc = Document(str(TITLE_TEMPLATE))
    ensure_heading_styles(doc)
    # the template has one table with a Тема row and a Выполнил row
    tbl = doc.tables[0]
    # Тема -> topic in the wide cell of row 0
    tcell = tbl.cell(0, len(tbl.columns) - 1)
    tcell.text = ""
    tr = tcell.paragraphs[0].add_run(TOPIC)
    tr.bold = True; tr.font.name = FONT; tr.font.size = Pt(14)
    # Выполнил -> student name. The "Выполнил" row is the last row; the
    # wide name cell is column index 4 (~7.6 cm); columns 3 and 5 are
    # thin spacers and column 6 is the signature box.
    last = tbl.rows[len(tbl.rows) - 1]
    scell = last.cells[4]
    scell.text = ""
    sr = scell.paragraphs[0].add_run(STUDENT)
    sr.bold = True; sr.font.name = FONT; sr.font.size = Pt(12)

    # margins
    for s in doc.sections:
        s.left_margin = Mm(25); s.right_margin = Mm(20)
        s.top_margin = Mm(20); s.bottom_margin = Mm(20)

    # The title-page template already ends with a section break (the
    # "Иннополис, 2026" line is the last content), so adding our own
    # WD_BREAK.PAGE here creates an extra empty page. Skip it; the
    # Оглавление heading naturally begins on the page that follows.

    # --- contents heading -----------------------------------------------
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run("Оглавление")
    cr.bold = True; cr.font.name = FONT; cr.font.size = Pt(16)
    cap.paragraph_format.space_after = Pt(6)
    # a real Word TOC field; it fills with entries + page numbers + dot
    # leaders when the field is updated in Word (no empty-list gap).
    toc_p = doc.add_paragraph()
    r = toc_p.add_run()
    fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-2" \h \z \u'
    fs = OxmlElement("w:fldChar"); fs.set(qn("w:fldCharType"), "separate")
    ph = OxmlElement("w:t")
    ph.text = "Обновите это поле, чтобы сформировать оглавление."
    fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end")
    for n in (fb, instr, fs, ph, fe):
        r._r.append(n)
    brk2 = doc.add_paragraph()
    brk2.add_run().add_break(WD_BREAK.PAGE)

    # --- body content ---------------------------------------------------
    # Chapter/section headings use real Heading styles so the TOC field
    # above populates with entries, page numbers, and dot leaders.
    # level 3 = a numbered reference-list entry (hanging indent).
    ref_n = 0
    for level, text in CONTENT:
        if level == 0:           # chapter heading
            p = doc.add_paragraph(text, style="Heading 1")
        elif level == 1:         # section heading
            p = doc.add_paragraph(text, style="Heading 2")
        elif level == 3:         # numbered reference entry
            ref_n += 1
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(f"{ref_n}. {text}")
            run.font.name = FONT; run.font.size = Pt(14)
        else:                    # body paragraph
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(1.25)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            run.font.name = FONT; run.font.size = Pt(14)

    add_page_numbers(doc)
    doc.save(str(OUT))
    print(f"written: {OUT}")


if __name__ == "__main__":
    main()
